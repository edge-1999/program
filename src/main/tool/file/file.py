import json
import os
from decimal import Decimal
import subprocess

import numpy as np
import pandas as pd
from docx import Document
import yaml


class FileOperations(object):
    def __init__(self, file_path):
        """文件进行操作"""
        self.file_path = None
        self.file_size = Decimal(0)
        self.identifying = False
        self.error_message = []
        try:
            if isinstance(file_path, str) and os.path.isfile(file_path):
                self.file_path = file_path
                self.identifying = True
            else:
                self.identifying = False
                self.error_message.append(f'当前文件 {self.file_path} 报错信息为：请输入正确的文件地址')
        except Exception as e:
            self.identifying = False
            self.error_message.append(f'执行__init__初始化时报错，当前文件 {self.file_path} 报错信息为：{e}')

    def def_file_json_load_manipulate(self):
        """json文件读取操作"""

        try:
            if len(self.file_path) > 4 and self.file_path[-5:].upper() == '.JSON':
                with open(self.file_path, 'r', encoding="utf-8") as f:
                    return json.load(f)
            self.identifying = False
            self.error_message.append(
                f'执行json文件读取操作时报错，当前文件 {self.file_path} 报错信息为：此文件不是json文件')
        except Exception as e:
            self.identifying = False
            self.error_message.append(f'执行json文件读取操作时报错，当前文件 {self.file_path} 报错信息为：{e}')

    def def_file_json_store_manipulate(self, data):
        """json文件写入操作"""
        try:
            if len(self.file_path) > 4 and self.file_path[-5:].upper() == '.JSON':
                with open(self.file_path, 'w', encoding="utf-8") as file_write:
                    file_write.write(json.dumps(dict(data, ensure_ascii=False)))
                    # json.dump(data, file_write, ensure_ascii=False)
                    self.identifying = True
                    print(f'数据覆盖写入文件 {self.file_path} 成功')
            self.identifying = False
            self.error_message.append(
                f'执行json文件写入操作时报错，当前文件 {self.file_path} 报错信息为：此文件不是json文件')
        except Exception as e:
            self.error_message.append(f'执行json文件写入操作时报错，当前文件 {self.file_path} 报错信息为：{e}')

    def def_file_xlsx_xls_search_for_manipulate(self, search_data):
        """xlsx_xls 搜索操作"""
        try:
            if len(self.file_path) > 4 and self.file_path[-5:].upper() == '.XLSX':
                xlsx = pd.ExcelFile(self.file_path)
                for name in xlsx.sheet_names:
                    data = pd.read_excel(xlsx, sheet_name=name, engine='openpyxl').fillna('')
                    for x in data.values:
                        for C in search_data:
                            if C in ''.join(str(x)):
                                self.identifying = True
                                return self.identifying
                                # yield self.file_path  # 方便日后做其余操作
            elif len(self.file_path) > 3 and self.file_path[-4:].upper() == '.XLS':
                xlsx = pd.ExcelFile(self.file_path)
                for name in xlsx.sheet_names:
                    data = pd.read_excel(xlsx, sheet_name=name).fillna('')
                    for x in data.values:
                        for C in search_data:
                            if C in ''.join(str(x)):
                                self.identifying = True
                                return self.identifying
                                # return self.identifying
                                # yield self.file_path  # 方便日后做其余操作
            self.identifying = False
            self.error_message.append(
                f'执行 xlsx_xls 文件搜索操作时报错，当前文件 {self.file_path} 报错信息为：此文件不是表格类型文件')
            return self.identifying
        except Exception as e:
            print('ok')
            self.identifying = False
            self.error_message.append(f'执行 xlsx_xls 文件搜索操作时报错，当前文件 {self.file_path} 报错信息为：{e}')

    def def_csv_search_for_manipulate(self, search_data):
        try:
            if len(self.file_path) > 3 and self.file_path[-4:].upper() == '.CSV':
                data_ = np.array(pd.read_csv(self.file_path).replace(np.nan, ''))
                for i in data_:
                    for j in search_data:
                        if j in i:
                            self.identifying = True
                            return self.identifying
                            # yield self.file_path
            self.identifying = False
            self.error_message.append(
                f'执行 csv 文件搜索操作时报错，当前文件 {self.file_path} 报错信息为：此文件不是 csv 类型文件')
        except Exception as e:
            self.identifying = False
            self.error_message.append(f'执行 csv 文件搜索操作时报错，当前文件 {self.file_path} 报错信息为：{e}')

    def def_docx_search_for_manipulate(self, search_data):
        """docx 文件搜索操作"""
        try:
            if len(self.file_path) > 4 and self.file_path[-5:].upper() == '.DOCX':
                doc = Document(self.file_path)
                docx_data = ""
                docx_data_tb = ""
                for para in doc.paragraphs:
                    docx_data += para.text + "\n"
                tbs = doc.tables
                for tb in tbs:
                    for row in tb.rows:
                        for cell in row.cells:
                            # docx_data_tb += f'{(cell.text):<5}'
                            docx_data_tb += '{({}):<5}'.format(cell.text)
                for dt in search_data:
                    if dt in docx_data or dt in docx_data_tb:
                        self.identifying = True
                        # return True
                        return self.identifying
                        # yield self.file_path
            self.identifying = False
            self.error_message.append(
                f'执行 docx 文件搜索操作时报错，当前文件 {self.file_path} 报错信息为：此文件不是 docx 类型文件')
        except Exception as e:
            self.identifying = False
            self.error_message.append(f'执行 docx 文件搜索操作时报错，当前文件 {self.file_path} 报错信息为：{e}')

    def def_txt_search_for_manipulate(self, search_data):
        try:
            if len(self.file_path) > 3 and self.file_path[-4:].upper() == '.TXT':
                txt_data = open(self.file_path)
                for dt in search_data:
                    if dt in txt_data:
                        self.identifying = True
                        return self.identifying
                        # yield self.file_path
            self.identifying = False
            self.error_message.append(
                f'执行 txt 文件搜索操作时报错，当前文件 {self.file_path} 报错信息为：此文件不是 txt 类型文件')
        except Exception as e:
            self.identifying = False
            self.error_message.append(f'执行 txt 文件搜索操作时报错，当前文件 {self.file_path} 报错信息为：{e}')

    def def_sql_search_for_manipulate(self, search_data):
        try:
            if len(self.file_path) > 3 and self.file_path[-4:].upper() == '.SQL':
                with open(self.file_path, 'r', encoding='utf', errors='ignore') as f:
                    for line in f:
                        for i in search_data:
                            if i in line:
                                self.identifying = True
                                return self.identifying
                                # yield self.file_path
            self.identifying = False
            self.error_message.append(
                f'执行 sql 文件搜索操作时报错，当前文件 {self.file_path} 报错信息为：此文件不是 sql 类型文件')
        except Exception as e:
            self.identifying = False
            self.error_message.append(f'执行 sql 文件搜索操作时报错，当前文件 {self.file_path} 报错信息为：{e}')

    def def_sh_search_for_manipulate(self, search_data):
        try:
            if len(self.file_path) > 3 and self.file_path[-4:].upper() == '.SH':
                with open(self.file_path, 'r', encoding='utf', errors='ignore') as f:
                    for line in f:
                        for i in search_data:
                            if i in line:
                                self.identifying = True
                                return self.identifying
                                # yield self.file_path
            self.identifying = False
            self.error_message.append(
                f'执行 sh 文件搜索操作时报错，当前文件 {self.file_path} 报错信息为：此文件不是 sh 类型文件')
        except Exception as e:
            self.identifying = False
            self.error_message.append(f'执行 sh 文件搜索操作时报错，当前文件 {self.file_path} 报错信息为：{e}')

    @property
    def def_file_filter_junk(self) -> bool:
        """判断文件是否是垃圾文件"""
        from src.config import file_path_name

        file_filter = file_path_name["filter_rubbish_file"] + file_path_name["Mac_filter_rubbish_file"]

        # file_name = os.path.dirname(self.file_path)
        file_name = os.path.basename(self.file_path)

        for file_filter_name in file_filter:
            if file_filter_name == file_name[: len(file_filter_name)]:
                self.identifying = True
                return self.identifying
        self.identifying = False
        return self.identifying

    @property
    def def_file_filter_junk_remove(self) -> bool:
        """删除垃圾文件"""
        if self.def_file_filter_junk:
            os.remove(self.file_path)
            self.identifying = True
            return self.identifying
        self.identifying = False
        return self.identifying

    def def_file_filter_appoint(self, file_data) -> bool:
        """判断文件是否是指定文件"""
        if isinstance(file_data, (list, set)) and file_data:
            for file_name in file_data:
                if os.path.basename(self.file_path) == os.path.basename(file_name):
                    self.identifying = True
            self.identifying = False
        elif isinstance(file_data, str):
            if os.path.basename(self.file_path) == os.path.basename(file_data):
                self.identifying = True
            self.identifying = False

        return self.identifying

    def def_file_filter_appoint_remove(self, file_data) -> bool:
        if self.def_file_filter_appoint(file_data):
            os.remove(self.file_path)
            self.identifying = True
            return self.identifying
        self.identifying = False
        return self.identifying

    @property
    def def_file_size(self):
        return Decimal(os.path.getsize(self.file_path))

    @property
    def def_file_output_size(self):
        """返回文件大小描述"""
        self.file_size = self.def_file_size
        file_size_default_radix = Decimal(1024)
        file_size_kb = self.file_size / file_size_default_radix
        if file_size_kb >= file_size_default_radix:
            file_size_mb = file_size_kb / file_size_default_radix
            if file_size_mb >= file_size_default_radix:
                file_size_gb = file_size_mb / file_size_default_radix
                if file_size_gb >= file_size_default_radix:
                    file_size_tb = file_size_gb / file_size_default_radix
                    return f'文件 {self.file_path} 大小为：{file_size_tb} TB'
                else:
                    return f'文件 {self.file_path} 大小为：{file_size_gb} GB'
            else:
                return f'文件 {self.file_path} 大小为：{file_size_mb} MB'
        else:
            return f'文件 {self.file_path} 大小为：{file_size_kb} KB'

    @property
    def def_file_read_encode(self):
        """读取整个文件内容"""
        try:
            self.identifying = True
            return open(self.file_path, 'r').read().encode()
        except Exception as e:
            self.identifying = False
            print(f'读取整个文件内容报错: {e}')

    def def_file_write(self, data) -> bool:
        decide = input("文件存在，是否覆盖，输入任意字符代表覆盖写入，不输入代表不覆盖: ")
        if decide:
            with open(self.file_path, 'a+', encoding='utf-8') as wa:
                wa.truncate(0)
                # 此处需要对data进行换行处理等
                wa.write(str(data))
                self.identifying = True
                return self.identifying
        else:
            self.identifying = False
            return self.identifying

    @property
    def def_file_encryption_lzma_base64_return_bytes(self) -> bytes:
        """
        把文件按照 lzma、base64 加密数据并返回
        """
        try:
            self.identifying = True
            import base64
            import lzma
            return base64.b64encode(lzma.compress(self.def_file_read_encode))
        except Exception as e:
            self.identifying = False
            print(f'把文件按照 lzma、base64 加密数据报错: {e}')
            return b''

    def def_file_write_lzma_base64(self, data) -> bool:
        """加密数据写入到文件"""
        decide = input("文件存在，是否覆盖，输入任意字符代表覆盖写入，不输入代表不覆盖: ")
        if decide:
            with open(self.file_path, 'a+', encoding='utf-8') as wa:
                wa.truncate(0)
                wa.write(f'import base64, lzma\nexec(lzma.decompress(base64.b64decode({data})))')
                self.identifying = True
                return self.identifying
        else:
            self.identifying = False
            return self.identifying


class FolderOperations(object):
    def __init__(self, folder_path):
        """对此文件夹进行操作"""
        self.identifying = False
        self.file_output_data = []
        self.error_message = []
        try:
            if isinstance(folder_path, str) and os.path.isdir(folder_path):
                self.folder_path = folder_path
                self.identifying = True
            else:
                self.error_message.append(f'当前文件夹 {self.folder_path} 报错信息为：请输入正确的文件地址')
        except Exception as e:
            self.error_message.append(f'当前文件夹 {self.folder_path} 报错信息为：{e}')

    def reveal_file_in_finder(self, file_path):
        subprocess.run(["open", "-R", file_path])

    def def_folder_search_content(self):
        """搜索文件夹下文件包含搜索内容的文件"""
        if self.identifying:
            from src.config import file_path_name
            while True:
                print(
                    f"目前识别文件格式为: {' '.join(i for i in file_path_name['file_search_for'])} 其他文件等待上线，"
                    f"输入条件（条件之间用一个空格隔开），按回车退出程序。", end='')
                search_content = input('')
                if search_content == '':
                    break
                # search_content = search_content.split()
                search_content = search_content.split(' ')
                set_path = set()
                for root, _, files in os.walk(self.folder_path):
                    if files:
                        for file_name in files:
                            file_manipulate = FileOperations(os.path.join(root, file_name))
                            self.identifying = True
                            if file_manipulate.def_file_size > 3 * 1024 * 1024:
                                decide = input(
                                    f"判断搜索文件:【{file_manipulate.file_path}】"
                                    f", 大小为【{file_manipulate.def_file_output_size}】"
                                    f", 请检查文件是否正常，直接回车跳过此文件搜索，输入任意按键并回车继续搜索"
                                    f">>> ")
                                if decide == '':
                                    self.identifying = False
                            if self.identifying and not file_manipulate.def_file_filter_junk:
                                if file_manipulate.file_path.lower()[-5:] == ".docx":
                                    if file_manipulate.def_docx_search_for_manipulate(search_content):
                                        set_path.add(file_manipulate.file_path)
                                elif file_manipulate.file_path.lower()[-5:] == ".xlsx" \
                                        or file_manipulate.file_path.lower()[-4:] == '.xls':
                                    if file_manipulate.def_file_xlsx_xls_search_for_manipulate(search_content):
                                        set_path.add(file_manipulate.file_path)
                                elif file_manipulate.file_path[-4:].lower() == ".txt":
                                    if file_manipulate.def_txt_search_for_manipulate(search_content):
                                        set_path.add(file_manipulate.file_path)
                                elif file_manipulate.file_path[-4:].lower() == ".csv":
                                    if file_manipulate.def_csv_search_for_manipulate(search_content):
                                        set_path.add(file_manipulate.file_path)
                                elif file_manipulate.file_path[-4:].lower() == ".sql":
                                    if file_manipulate.def_sql_search_for_manipulate(search_content):
                                        set_path.add(file_manipulate.file_path)
                                elif file_manipulate.file_path[-3:].lower() == ".sh":
                                    if file_manipulate.def_sh_search_for_manipulate(search_content):
                                        set_path.add(file_manipulate.file_path)

                if len(set_path):
                    print('\n>>> 一共{}文件如下所示 >>> \n\t{}'.format(
                        len(set_path), '\n\t'.join(file_name for file_name in set_path)))
                    self.identifying = True
                    while self.identifying:
                        determining_the_folder_where_the_file_is_located = input('是否打开文件 y|Y ，直接回车跳过>>> ')
                        if determining_the_folder_where_the_file_is_located:
                            if determining_the_folder_where_the_file_is_located[0].upper() == 'Y':
                                file_path_open_ = input('请输入打开文件的路径')
                                if file_path_open_ in set_path:
                                    self.reveal_file_in_finder(file_path_open_)
                                else:
                                    print('请输入展示中的文件路径 >>> \n\t{}'.format(
                                        '\n\t'.join(file_name for file_name in set_path)))
                        else:
                            self.identifying = False

                else:
                    print(f"\n在路径【{self.folder_path}】中未找到含有【{' '.join(search_content)}】内容的文件")

    @property
    def def_folder_obtain_full_path(self):
        """获取此文件夹目录的所有文件和文件夹"""
        for root, dirs, files in os.walk(self.folder_path):
            for file_name in files:
                file_manipulate = FileOperations(os.path.join(root, file_name))
                if file_manipulate.def_file_filter_junk:
                    self.file_output_data.append(
                        [file_manipulate.file_path, '文件', '无效', file_name])
                self.file_output_data.append(
                    [file_manipulate.file_path, '文件', '有效', file_name])

            for dir_name in dirs:
                dir_root_path = os.path.join(root, dir_name)
                if '.' == dir_name[:1]:
                    self.file_output_data.append(
                        [dir_root_path, '文件夹', '无效', dir_name])
                self.file_output_data.append(
                    [dir_root_path, '文件夹', '有效', dir_name])
            self.identifying = True
        return self.identifying

    def def_folder_obtain_full_path_print(self, path=None, depth=0, site=None):
        """递归获取目录层级"""
        if site is None and path is None:
            site = []
            path = self.folder_path

        void_num = 0
        filenames_list = os.listdir(path)

        for item in filenames_list:
            if '._' == item[:2]:
                continue
            string_list = ["│   " for _ in range(depth - void_num - len(site))]
            for s in site:
                string_list.insert(s, "    ")

            if item != filenames_list[-1]:
                string_list.append("├── ")
            else:
                # 本级目录最后一个文件：转折处
                string_list.append("└── ")
                void_num += 1
                # 添加当前已出现转折的层级数
                site.append(depth)
            print("".join(string_list) + item)

            new_item = path + '/' + item
            if os.path.isdir(new_item):
                self.def_folder_obtain_full_path_print(new_item, depth + 1, site)
            if item == filenames_list[-1]:
                void_num -= 1
                # 移除当前已出现转折的层级数
                site.pop()

    def def_folder_obtain_full_path_save_xlsx(self, xlsx_path):
        """获取文件夹下所有的文件、目录并输出到 xlsx 文件 """
        if self.def_folder_obtain_full_path:
            new_pd_file = pd.DataFrame(self.file_output_data)
            new_pd_file.to_excel(xlsx_path, index=False, header=['路径', '类型', '是否有效', '名称'])
            self.file_output_data = []
            self.identifying = True
        return self.identifying

    def def_folder_remove_file_junk(self):
        """删除文件夹内的垃圾文件"""
        for root, _, files in os.walk(self.folder_path):
            if files:
                for file_name in files:
                    file_m = FileOperations(os.path.join(root, file_name))
                    if file_m.identifying:
                        if file_m.def_file_filter_junk_remove:
                            print(f'删除 {file_m.file_path} 文件成功！')

    def def_folder_search_for_file_folder_names_input(self):
        """查找文件夹下文件或文件夹名称包含[]中的文件或文件夹"""
        while True:
            print("请输入要搜索的名称，输入条件（条件之间用一个空格隔开），按回车退出程序。", end='')
            search_content = input('')
            if search_content == '':
                break
            # search_content = search_content.split()
            search_content = search_content.split(' ')
            set_path = set()
            for root, dirs, files in os.walk(self.folder_path):
                if files:
                    for file_name in files:
                        for filter_content in search_content:
                            if filter_content in file_name:
                                set_path.add(os.path.join(root, file_name))
                if dirs:
                    for dir_name in dirs:
                        for filter_content in search_content:
                            if filter_content in dir_name:
                                set_path.add(os.path.join(root, dir_name))
            if len(set_path):
                print(f'\n>>> 一共{len(set_path)}文件或文件夹')
                for file_name in set_path:
                    print(f'【{file_name}】')
            else:
                print(f"\n在路径【{self.folder_path}】中未找到含有【{' '.join(search_content)}】内容名称的文件或文件夹")


def generate_file_tree_local(path: str, depth: int, site: list):
    """
    递归打印文件目录树状图（使用局部变量）

    :param path: 根目录路径
    :param depth: 根目录、文件所在的层级号
    :param site: 存储出现转折的层级号
    :return: None
    """
    void_num = 0
    filenames_list = os.listdir(path)

    for item in filenames_list:
        if '._' == item[:2]:
            continue
        string_list = ["│   " for _ in range(depth - void_num - len(site))]
        for s in site:
            string_list.insert(s, "    ")

        if item != filenames_list[-1]:
            string_list.append("├── ")
        else:
            # 本级目录最后一个文件：转折处
            string_list.append("└── ")
            void_num += 1
            # 添加当前已出现转折的层级数
            site.append(depth)
        print("".join(string_list) + item)

        new_item = path + '/' + item
        if os.path.isdir(new_item):
            generate_file_tree_local(new_item, depth + 1, site)
        if item == filenames_list[-1]:
            void_num -= 1
            # 移除当前已出现转折的层级数
            site.pop()


class FileTree(object):
    """
    st = FileTree('/Users/li/Downloads/尚硅谷Docker实战教程（docker教程天花板）', '/Users/li/Downloads/工作簿1_out.xlsx').run()
    """

    def __init__(self, directory_path, to_excel_path):
        """
        directory_path:目录例如：'/Users/li/Downloads/2023.0619集团公司总部内控手册'
        to_excel_path:保存为excel_path的路径 例如：'/Users/li/Downloads/工作簿1_out.xlsx'
        header_title：标题 例如：["类型", "文件路径", "文件目录", "父级目录"]
        """
        self.directory_path = directory_path
        self.to_excel_path = to_excel_path
        self.data = []  # 替换为 np 或 pd
        self.status = False

    def _gain_directory(self, directory_path):
        for root, dirs, files in os.walk(directory_path):
            for file in files:
                if '._' == file[:2] or '.DS_Store' == file:
                    continue
                p1 = str(root + '/' + file).replace(self.directory_path, '')
                p2 = root.replace(self.directory_path, '')
                pd_data = [
                    "文件", p1.replace('/', '', 1) if '/' == p1[:1] else p1,
                    file, p2.replace('/', '', 1) if '/' == p2[:1] else p2]
                print(pd_data)
                self.data.append(pd_data)
            for dir_ in dirs:
                if '._' == dir_[:2]:
                    continue
                p1 = str(root + '/' + dir_).replace(self.directory_path, '')
                p2 = root.replace(self.directory_path, '')
                pd_data = [
                    "目录", p1.replace('/', '', 1) if '/' == p1[:1] else p1, dir_,
                    p2.replace('/', '', 1) if '/' == p2[:1] else p2]
                print(pd_data)
                self.data.append(pd_data)
        # return self.data

    def _to_excel_file(self, data, to_excel_path):
        pd.DataFrame(data).to_excel(to_excel_path, index=False, header=["类型", "文件路径", "文件目录", "父级目录"])
        self.status = True
        return self.status

    def run(self):
        self._gain_directory(self.directory_path)
        return self._to_excel_file(np.array(self.data), self.to_excel_path)


def _gain_directory(directory_path: str, content: list):
    """
    输入需要查找的地址：/Volumes/LaCie/MyWork
    输入需要查找的内容：["科研"]
    直接执行：_gain_directory('/Users/li/Downloads', ['k8s'])
    """
    for root, dirs, files in os.walk(directory_path):
        for file in files:
            for i in content:
                if i in file:
                    print(f'{root}/{file}')

        for dir_ in dirs:
            for i in content:
                if i in dir_:
                    print(f'{root}/{dir_}')


def yaml_generator(file_path):
    """生成器读取yaml文件"""
    with open(file_path, 'r') as file:
        for line in file:
            yield yaml.safe_load(line)


class YamlIterator:
    """迭代器读取yaml文件"""

    def __init__(self, file_path):
        self.file_path = file_path

    def __iter__(self):
        with open(self.file_path, 'r') as file:
            for line in file:
                yield yaml.safe_load(line)


def yaml_generator_load(file_path):
    """读取yaml文件"""
    with open(file_path, 'r') as yaml_file:
        return yaml.safe_load(yaml_file)


def yaml_generator_wri(file_path, data):
    """写入yaml文件"""
    with open(file_path, 'w') as yaml_file:
        yaml.dump(data, yaml_file)
